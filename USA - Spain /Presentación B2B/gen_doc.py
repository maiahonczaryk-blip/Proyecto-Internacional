#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
gen_doc.py — v2.2
Genera el Guión Completo del Webinario B2B — Beyond Borders
RE/MAX Inmomás International Partner Program
Diálogos EN inglés + 🇪🇸 español (misma extensión y detalle)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Colores corporativos ─────────────────────────────────────────────────────
RED   = RGBColor(0xC8, 0x10, 0x2E)
GREY  = RGBColor(0x80, 0x80, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK  = RGBColor(0x1A, 0x1A, 0x2E)
BLUE  = RGBColor(0x22, 0x66, 0xAA)
DKGRN = RGBColor(0x00, 0x70, 0xC0)

# ─── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def set_table_borders(table):
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    borders = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'D0D0D0')
        borders.append(el)
    tblPr.append(borders)


def add_paragraph(doc, text='', bold=False, italic=False,
                  color=None, size=11, align=WD_ALIGN_PARAGRAPH.LEFT,
                  indent_left=None, space_before=None, space_after=None):
    p  = doc.add_paragraph()
    p.alignment = align
    pPr = p.paragraph_format
    if indent_left  is not None: pPr.left_indent  = Cm(indent_left)
    if space_before is not None: pPr.space_before = Pt(space_before)
    if space_after  is not None: pPr.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold       = bold
        run.italic     = italic
        run.font.size  = Pt(size)
        run.font.color.rgb = color if color else BLACK
    return p


def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 1']
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size      = Pt(18)
    run.font.color.rgb = RED
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 2']
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size      = Pt(14)
    run.font.color.rgb = DARK
    return p


def add_script(doc, text, speaker=None):
    """Guión en inglés — cursiva con sangría."""
    if speaker:
        sp = doc.add_paragraph()
        sp.paragraph_format.left_indent  = Cm(1)
        sp.paragraph_format.space_before = Pt(6)
        sp.paragraph_format.space_after  = Pt(2)
        run = sp.add_run(speaker)
        run.bold = True
        run.font.size      = Pt(10)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.italic = True
    run.font.size      = Pt(10.5)
    run.font.color.rgb = BLACK
    return p


def add_script_es(doc, text, speaker=None):
    """Versión española — misma extensión y detalle que el inglés."""
    # Etiqueta de idioma
    lbl = doc.add_paragraph()
    lbl.paragraph_format.left_indent  = Cm(1)
    lbl.paragraph_format.space_before = Pt(5)
    lbl.paragraph_format.space_after  = Pt(2)
    r = lbl.add_run('🇪🇸  Versión en español:')
    r.bold = True
    r.font.size      = Pt(9.5)
    r.font.color.rgb = BLUE

    if speaker:
        sp = doc.add_paragraph()
        sp.paragraph_format.left_indent  = Cm(1)
        sp.paragraph_format.space_before = Pt(2)
        sp.paragraph_format.space_after  = Pt(2)
        run = sp.add_run(speaker)
        run.bold = True
        run.font.size      = Pt(10)
        run.font.color.rgb = RGBColor(0x22, 0x44, 0x88)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(8)
    run = p.add_run(text)
    run.italic = True
    run.font.size      = Pt(10.5)
    run.font.color.rgb = RGBColor(0x11, 0x11, 0x44)
    return p


def add_instruction(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.italic = True
    run.font.size      = Pt(9.5)
    run.font.color.rgb = GREY
    return p


def add_separator(doc):
    p = doc.add_paragraph('─' * 70)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        run.font.size = Pt(8)


def add_footer(doc):
    section = doc.sections[0]
    footer  = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.clear()
    run = fp.add_run('RE/MAX Inmomás International · Uso interno · 2026')
    run.font.size      = Pt(8)
    run.font.color.rgb = GREY
    run.italic = True


# ─── Construcción del documento ───────────────────────────────────────────────

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

add_footer(doc)

# ══════════════════════════════════════════════════════════════════════════════
# PORTADA / TÍTULO
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after  = Pt(4)
run = p.add_run('GUIÓN COMPLETO DEL WEBINARIO B2B — Beyond Borders')
run.bold = True
run.font.size      = Pt(22)
run.font.color.rgb = RED

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(2)
run2 = p2.add_run(
    'RE/MAX Inmomás International Partner Program\n'
    'Dirigido a Realtors de EE.UU. y Canadá'
)
run2.font.size      = Pt(13)
run2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(2)
run3 = p3.add_run('"Beyond Borders: The North American Realtor\'s Blueprint to Spain\'s Expat Boom"')
run3.italic = True
run3.font.size      = Pt(11)
run3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.paragraph_format.space_after = Pt(16)
run4 = p4.add_run('Duración total: 48–52 minutos | Formato: Bilingüe Inglés / Español')
run4.font.size      = Pt(10)
run4.font.color.rgb = GREY

add_separator(doc)

# ══════════════════════════════════════════════════════════════════════════════
# FICHA TÉCNICA
# ══════════════════════════════════════════════════════════════════════════════
add_paragraph(doc, '📋  FICHA TÉCNICA / FACT SHEET',
              bold=True, color=RED, size=13, space_before=10, space_after=6)

ficha_data = [
    ('Título',             '"Beyond Borders: The North American Realtor\'s Blueprint to Spain\'s Expat Boom"'),
    ('Subtítulo',          '"How to secure your clients\' relocation to Europe and capture premium 25% referral fees — without leaving your desk."'),
    ('Duración',           '48 – 52 minutos + Q&A'),
    ('Formato',            'Webinario en vivo (Zoom / Teams) con grabación'),
    ('Presentadores',      '2–3 Agentes de RE/MAX Inmomás (americanos/canadienses viviendo en España)'),
    ('Partners invitados', 'Fuster & Associates (Legal) · UCI (Financiación) · Inmomás Holidays (Rentas vacionales)'),
    ('Audiencia objetivo', 'Realtors y Broker Associates de EE.UU. y Canadá'),
    ('Herramienta visual', 'Presentación de 15 diapositivas (PowerPoint / Canva)'),
]

t1 = doc.add_table(rows=1 + len(ficha_data), cols=2)
t1.alignment = WD_TABLE_ALIGNMENT.LEFT
t1.style = 'Table Grid'
set_table_borders(t1)

hdr = t1.rows[0].cells
set_cell_bg(hdr[0], 'C8102E')
set_cell_bg(hdr[1], 'C8102E')
for ci, txt in enumerate(('Campo', 'Detalle')):
    run = hdr[ci].paragraphs[0].add_run(txt)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)
    hdr[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for i, (campo, detalle) in enumerate(ficha_data, start=1):
    row = t1.rows[i].cells
    if i % 2 == 0:
        set_cell_bg(row[0], 'FFF0F0')
        set_cell_bg(row[1], 'FFF0F0')
    run_c = row[0].paragraphs[0].add_run(campo)
    run_c.bold = True
    run_c.font.size      = Pt(9.5)
    run_c.font.color.rgb = RED
    run_d = row[1].paragraphs[0].add_run(detalle)
    run_d.font.size      = Pt(9.5)
    run_d.font.color.rgb = BLACK

t1.columns[0].width = Cm(4.5)
t1.columns[1].width = Cm(12.0)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# CRONOGRAMA
# ══════════════════════════════════════════════════════════════════════════════
add_paragraph(doc, '⏱️  CRONOGRAMA', bold=True, color=RED, size=13,
              space_before=10, space_after=6)

crono_data = [
    ('1', 'Apertura, Bienvenida y Presentación',    '00:00 – 08:00', 'Agentes RE/MAX Inmomás'),
    ('2', 'El Gran Éxodo Norteamericano (Datos)',    '08:00 – 20:00', 'Agentes RE/MAX Inmomás'),
    ('3', 'El Ecosistema 360° – Partners invitados', '20:00 – 32:00', 'Partners (1.5–2 min c/u)'),
    ('4', 'La Fórmula Financiera: Tu 25%',           '32:00 – 40:00', 'Agentes RE/MAX Inmomás'),
    ('5', 'El Portal, CTA y Q&A en vivo',            '40:00 – 52:00', 'Agentes RE/MAX Inmomás'),
]

t2 = doc.add_table(rows=1 + len(crono_data), cols=4)
t2.alignment = WD_TABLE_ALIGNMENT.LEFT
t2.style = 'Table Grid'
set_table_borders(t2)

crono_hdrs = ('Bloque', 'Tema', 'Minutos', 'Presentador')
for ci, h in enumerate(crono_hdrs):
    c = t2.rows[0].cells[ci]
    set_cell_bg(c, 'C8102E')
    run = c.paragraphs[0].add_run(h)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for i, row_data in enumerate(crono_data, start=1):
    row = t2.rows[i].cells
    if i % 2 == 0:
        for c in row: set_cell_bg(c, 'FFF5F5')
    for ci, txt in enumerate(row_data):
        run = row[ci].paragraphs[0].add_run(txt)
        run.font.size = Pt(9.5)
        if ci == 0:
            run.bold = True
            run.font.color.rgb = RED
        else:
            run.font.color.rgb = BLACK
        if ci in (0, 2):
            row[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

t2.columns[0].width = Cm(1.5)
t2.columns[1].width = Cm(7.5)
t2.columns[2].width = Cm(3.0)
t2.columns[3].width = Cm(4.5)

doc.add_paragraph()
add_separator(doc)

# ══════════════════════════════════════════════════════════════════════════════
# BLOQUE 1 — APERTURA Y BIENVENIDA
# ══════════════════════════════════════════════════════════════════════════════
add_heading1(doc, 'BLOQUE 1 — APERTURA Y BIENVENIDA')
add_paragraph(doc, '⏰ 00:00 – 08:00 | Presentan: Agentes de RE/MAX Inmomás',
              italic=True, color=GREY, size=10)

# ── SLIDE 1 ──────────────────────────────────────────────────────────────────
add_heading2(doc, '🎬 SLIDE 1: PORTADA — "Beyond Borders"')
add_instruction(doc,
    '[Acción de producción: La pantalla muestra la portada con música de fondo '
    'suave. El presentador entra en cámara sonriendo.]')

add_script(doc,
    '"Hello, everyone, and welcome! I\'m so glad you\'re here today. My name is [Nombre], '
    'and on behalf of the entire RE/MAX Inmomás International team, I want to genuinely '
    'thank you for carving out this hour from your busy schedule.\n\n'
    'I know that as a Realtor, your time is your most valuable asset. So let me make you '
    'a promise right now: by the end of this session, you\'re going to walk away with a '
    'completely new revenue stream that requires zero new licensing, zero new market '
    'knowledge, and zero extra hours of driving clients around town.\n\n'
    'This is a Realtor-to-Realtor conversation. No sales pitch. No corporate speak. Just '
    'two colleagues who found something remarkable — and want to share it with you."',
    speaker='GUIÓN EN INGLÉS (Agente 1 — voz principal de apertura):')

add_script_es(doc,
    '"¡Hola a todos, y bienvenidos! Me alegra muchísimo que estéis aquí hoy. Mi nombre es [Nombre], '
    'y en nombre de todo el equipo de RE/MAX Inmomás International, quiero agradeceros sinceramente '
    'que hayáis reservado esta hora en vuestra apretada agenda.\n\n'
    'Sé que, como Realtors, vuestro tiempo es vuestro activo más valioso. Así que permitidme '
    'haceros una promesa ahora mismo: al final de esta sesión, vais a salir de aquí con una '
    'fuente de ingresos completamente nueva que no requiere ninguna licencia adicional, ningún '
    'conocimiento de mercado nuevo, y cero horas extra llevando clientes de propiedad en propiedad.\n\n'
    'Esto es una conversación entre colegas inmobiliarios. Sin discurso de ventas. Sin jerga '
    'corporativa. Solo dos compañeros que han encontrado algo extraordinario — y quieren '
    'compartirlo con vosotros."',
    speaker='GUIÓN EN ESPAÑOL (Agente 1 — voz principal de apertura):')

# ── SLIDE 2 ──────────────────────────────────────────────────────────────────
add_heading2(doc, '🎬 SLIDE 2: PRESENTACIÓN DE LOS PONENTES — "Meet Your Hosts"')
add_instruction(doc,
    '[Diapositiva con fotos profesionales de los agentes, sus banderas de origen '
    'y de España, nombre y cargo.]')

add_script(doc,
    '"Before we get into the big numbers, let me tell you who we are — because I think '
    'that\'s actually the most important part.\n\n'
    'My name is [Nombre completo]. I spent [X] years as a licensed Realtor in '
    '[California / Florida / Texas / etc.]. I worked with buyers, sellers, investors — '
    'I know the grind, I know the paperwork, I know what it takes to close a deal in a '
    'competitive market.\n\n'
    'Three years ago, I made the move to Spain. And I want to be honest with you: it '
    'wasn\'t a retirement plan, it was the best business decision of my life. Today I '
    'specialize in helping North American buyers find their perfect property on the Costa '
    'Blanca — and I do it under the RE/MAX brand, the most trusted name in international '
    'real estate."',
    speaker='GUIÓN EN INGLÉS (Agente 1 — se presenta):')

add_script_es(doc,
    '"Antes de entrar en los grandes números, dejadme contaros quiénes somos — porque creo '
    'que ese es, en realidad, el punto más importante.\n\n'
    'Mi nombre es [Nombre completo]. Trabajé [X] años como Realtor con licencia en '
    '[California / Florida / Texas / etc.]. He trabajado con compradores, vendedores, '
    'inversores — conozco el esfuerzo diario, conozco el papeleo, sé lo que se necesita '
    'para cerrar una operación en un mercado competitivo.\n\n'
    'Hace tres años, di el salto a España. Y quiero ser honesto con vosotros: no fue un '
    'plan de jubilación anticipada, fue la mejor decisión de negocio de mi vida. Hoy me '
    'especializo en ayudar a compradores norteamericanos a encontrar su propiedad perfecta '
    'en la Costa Blanca — y lo hago bajo la marca RE/MAX, el nombre más reconocido del '
    'sector inmobiliario internacional."',
    speaker='GUIÓN EN ESPAÑOL (Agente 1 — se presenta):')

add_script(doc,
    '"And I\'m [Nombre completo]. I was a Broker Associate in [Toronto / Vancouver / '
    'Montreal] for [X] years. My clients were high-net-worth buyers who knew what they '
    'wanted. The thing that finally pushed me to Spain was a combination of the cost of '
    'living, the lifestyle, and — quite honestly — those Canadian winters.\n\n'
    'Today, my spouse and I live in [Alicante / Valencia / Murcia], our kids go to an '
    'international school, and I spend my working hours connecting people with '
    'life-changing opportunities in one of the most beautiful places on earth.\n\n'
    'We are not corporate representatives. We are you — Realtors who made the leap. And '
    'that\'s exactly why we can help your clients make it too."',
    speaker='GUIÓN EN INGLÉS (Agente 2 — se presenta):')

add_script_es(doc,
    '"Y yo soy [Nombre completo]. Fui Broker Associate en [Toronto / Vancouver / '
    'Montreal] durante [X] años. Mis clientes eran compradores de alto poder adquisitivo '
    'que sabían exactamente lo que querían. Lo que finalmente me impulsó a dar el paso a '
    'España fue una combinación del coste de vida, el estilo de vida y — siendo completamente '
    'honesto — esos inviernos canadienses que no tienen fin.\n\n'
    'Hoy, mi pareja y yo vivimos en [Alicante / Valencia / Murcia], nuestros hijos van a '
    'un colegio internacional, y dedico mis horas de trabajo a conectar a personas con '
    'oportunidades que cambian sus vidas en uno de los lugares más hermosos del mundo.\n\n'
    'No somos representantes corporativos. Somos vosotros — Realtors que se atrevieron '
    'a dar el salto. Y es exactamente por eso que podemos ayudar a que vuestros clientes '
    'también lo den."',
    speaker='GUIÓN EN ESPAÑOL (Agente 2 — se presenta):')

add_instruction(doc, '[Pausa de 5 segundos. El tono se vuelve cálido y cómplice.]')
add_script(doc,
    '"We\'ve been in your exact position: watching clients talk about moving abroad, not '
    'knowing how to help them, and quietly losing that relationship. Today we\'re going '
    'to fix that. Forever."')
add_script_es(doc,
    '"Hemos estado exactamente en vuestra posición: escuchando a clientes hablar de mudarse '
    'al extranjero, sin saber cómo ayudarles, y perdiéndoles sin hacer ruido. Hoy vamos '
    'a solucionar eso. Para siempre."')

# ── SLIDE 3 ──────────────────────────────────────────────────────────────────
add_heading2(doc, '🎬 SLIDE 3: TESTIMONIO DE VIDA — "Why We Made the Leap"')
add_instruction(doc,
    '[Fotos personales: familia en la playa, terraza mediterránea, oficina con vistas '
    'al mar, mercado local.]')

add_script(doc,
    '"Let me take 60 seconds to paint you a picture of my typical Tuesday.\n\n'
    'I wake up at 7:30 AM. There\'s sunlight coming through the window. I have breakfast '
    'on my terrace with a view of the Mediterranean. By 9 AM I\'m at the office — and I '
    'mean a proper office, with a team, with systems, with the full RE/MAX infrastructure '
    'behind me. I spend my morning on calls with clients in California, Texas, New York. '
    'By 1 PM I\'m done with most of my digital meetings.\n\n'
    'Then I drive 8 minutes to the beach for lunch.\n\n'
    'That is the life I\'m selling. And the thing is — it\'s my actual life. I don\'t '
    'need to exaggerate, I just need to show it."',
    speaker='GUIÓN EN INGLÉS (Agente 1 — testimonio emocional):')

add_script_es(doc,
    '"Dejadme 60 segundos para pintaros una imagen de mi martes típico.\n\n'
    'Me despierto a las 7:30 de la mañana. Entra luz del sol por la ventana. Desayuno '
    'en mi terraza con vistas al Mediterráneo. A las 9 estoy en la oficina — y hablo de '
    'una oficina de verdad, con equipo, con sistemas, con toda la infraestructura de '
    'RE/MAX detrás de mí. Paso la mañana en llamadas con clientes en California, Texas, '
    'Nueva York. A la 1 del mediodía ya he terminado la mayoría de mis reuniones digitales.\n\n'
    'Después conduzco 8 minutos hasta la playa para comer.\n\n'
    'Esa es la vida que vendo. Y lo importante es que — es mi vida real. No necesito '
    'exagerar nada. Solo necesito mostrarla."',
    speaker='GUIÓN EN ESPAÑOL (Agente 1 — testimonio emocional):')

add_script(doc,
    '"For me, the tipping point was healthcare. In Canada, I paid a fortune in taxes and '
    'still waited months for specialist appointments. Here, my entire family has '
    'world-class private healthcare coverage for less than $200 a month.\n\n'
    'My clients from Canada always ask me: \'Is it really that good?\' And I show them '
    'my utility bills, my grocery receipts, my kids\' school grades — and then I invite '
    'them to come see it for themselves. The VIP tour we organize does the closing. Not '
    'me. Spain does the closing."',
    speaker='GUIÓN EN INGLÉS (Agente 2 — testimonio canadiense):')

add_script_es(doc,
    '"Para mí, el punto de inflexión fue la sanidad. En Canadá, pagaba una fortuna en '
    'impuestos y aun así esperaba meses para ver a un especialista. Aquí, toda mi familia '
    'tiene cobertura sanitaria privada de primer nivel por menos de 200 dólares al mes.\n\n'
    'Mis clientes canadienses siempre me preguntan: "¿De verdad es tan bueno?" Y les '
    'enseño mis facturas de suministros, mis tickets del supermercado, las notas de mis '
    'hijos en el colegio — y luego los invito a venir a verlo por sí mismos. El VIP Tour '
    'que organizamos es el que cierra la venta. No yo. España cierra la venta."',
    speaker='GUIÓN EN ESPAÑOL (Agente 2 — testimonio canadiense):')

# ── SLIDE 4 ──────────────────────────────────────────────────────────────────
add_heading2(doc, '🎬 SLIDE 4: LA AGENDA — "Today\'s Playbook"')

add_script(doc,
    '"Here\'s what we\'re going to cover today — and I promise we\'ll stay tight on time '
    'because I know you have clients to call.\n\n'
    'First, we\'ll show you the data — the migration numbers that prove this isn\'t a '
    'trend, it\'s a structural shift. Second, you\'re going to meet our full ecosystem of '
    'partners: a legal team, a mortgage specialist, and a vacation rental management company '
    '— because we don\'t just sell houses, we deliver a complete transition service. Third, '
    'we\'ll break down the commission math — exactly how your 25% works and what it looks '
    'like on a real deal. And finally, we\'ll do a live demo of our partner portal so you can '
    'see exactly how simple it is to send us a referral.\n\n'
    'Any questions along the way, drop them in the chat — we\'ll address everything in '
    'the Q&A at the end.\n\n'
    'Let\'s get started."',
    speaker='GUIÓN EN INGLÉS:')

add_script_es(doc,
    '"Esto es lo que vamos a cubrir hoy — y os prometo que respetaremos el tiempo porque '
    'sé que tenéis clientes a los que llamar.\n\n'
    'Primero, os mostraremos los datos — las cifras de migración que demuestran que esto '
    'no es una tendencia pasajera, sino un cambio estructural. Segundo, conoceréis nuestro '
    'ecosistema completo de partners: un equipo legal, un especialista en hipotecas y una '
    'empresa de gestión de alquileres vacacionales — porque no solo vendemos casas, '
    'ofrecemos un servicio de transición completo. Tercero, desglosaremos la fórmula de '
    'la comisión — exactamente cómo funciona vuestro 25% y cómo queda en una operación real. '
    'Y por último, haremos una demo en vivo del portal de partners para que veáis lo sencillo '
    'que es enviarnos una referencia.\n\n'
    'Cualquier pregunta durante la sesión, déjadla en el chat — lo resolveremos todo en '
    'el Q&A del final.\n\n'
    '¡Empezamos!"',
    speaker='GUIÓN EN ESPAÑOL:')

add_separator(doc)

# ══════════════════════════════════════════════════════════════════════════════
# BLOQUE 2 — EL GRAN ÉXODO NORTEAMERICANO
# ══════════════════════════════════════════════════════════════════════════════
add_heading1(doc, 'BLOQUE 2 — EL GRAN ÉXODO NORTEAMERICANO')
add_paragraph(doc, '⏰ 08:00 – 20:00 | Presentan: Agentes de RE/MAX Inmomás',
              italic=True, color=GREY, size=10)

# ── SLIDE 5 ──────────────────────────────────────────────────────────────────
add_heading2(doc, '🎬 SLIDE 5: "The North American Exodus" — Los Números que Cambian Todo')
add_instruction(doc,
    '[Diapositiva de alto impacto: mapa animado de EE.UU. y Canadá → España, con cifras grandes.]')

add_script(doc,
    '"Let me ask you something. In the last 12 months, how many of your clients have '
    'mentioned thinking about moving abroad? Maybe to Mexico, Costa Rica, Portugal, Spain?\n\n'
    'Go ahead and drop a number in the chat right now. I\'m genuinely curious.\n\n'
    '[Pausa de 15 segundos para respuestas del chat.]\n\n'
    'Right. Every single one of you has had that conversation. And here\'s the problem: '
    'most of us, when that client says \'I\'m thinking about Spain,\' we smile, we say '
    '\'wow, how exciting\' — and then we lose them. Because we don\'t have the '
    'infrastructure to help them.\n\n'
    'Until today.\n\n'
    'Now look at these numbers."',
    speaker='GUIÓN EN INGLÉS:')

add_script_es(doc,
    '"Dejadme haceros una pregunta. En los últimos 12 meses, ¿cuántos de vuestros clientes '
    'han mencionado que están pensando en mudarse al extranjero? ¿Quizás a México, Costa Rica, '
    'Portugal, España?\n\n'
    'Escribid un número en el chat ahora mismo. Tengo genuina curiosidad.\n\n'
    '[Pausa de 15 segundos para respuestas del chat.]\n\n'
    'Exacto. Todos y cada uno de vosotros habéis tenido esa conversación. Y aquí está el '
    'problema: la mayoría de nosotros, cuando ese cliente dice "estoy pensando en España", '
    'sonreímos, decimos "¡qué emocionante!" — y luego lo perdemos. Porque no tenemos la '
    'infraestructura para ayudarle.\n\n'
    'Hasta hoy.\n\n'
    'Ahora mirad estas cifras."',
    speaker='GUIÓN EN ESPAÑOL:')

add_instruction(doc, '[Transición a estadísticas — leer cada una con énfasis dramático:]')

stats = [
    ('🔢', 'Over 40 million Americans have considered relocating abroad — that\'s not a fringe idea, '
           'that\'s a movement. [Fuente: Gallup / State Department]'),
    ('🔢', 'Spain received a record 3.9 million international residents in 2023, with North Americans '
           'among the fastest-growing segments. [Fuente: INE España]'),
    ('🔢', 'Searches for \'move to Spain\' and \'buy property in Spain\' from the US have increased '
           'by 340% in the last three years. [Fuente: Google Trends]'),
    ('🔢', 'For Canadians specifically, Spain\'s digital nomad visa and golden visa program have '
           'triggered a 35% surge in relocation inquiries since 2022. [Fuente: Fuster & Associates / UCI]'),
    ('🔢', 'The Costa Blanca alone — our market — saw over €2.1 billion in international real estate '
           'transactions in 2023, with North Americans in the top 5 buyer nationalities.'),
]
for icon, stat in stats:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f'{icon}  {stat}')
    run.italic = True
    run.font.size      = Pt(10.5)
    run.font.color.rgb = BLACK

add_script(doc,
    '"These are not projections. These are not marketing numbers. These are documented, '
    'official statistics. The demand is here. The question is: who is going to capture it?"')
add_script_es(doc,
    '"Estas no son proyecciones. No son cifras de marketing. Son estadísticas oficiales '
    'documentadas. La demanda existe. La pregunta es: ¿quién va a captarla?"')

# ── SLIDE 6 ──────────────────────────────────────────────────────────────────
add_heading2(doc, '🎬 SLIDE 6: "Why Spain? Why NOW?" — Los Motivos que Impulsan la Migración')
add_instruction(doc,
    '[Infografía comparativa: EE.UU./Canadá vs España en costo de vida, salud, seguridad, clima.]')

add_script(doc,
    '"Now let\'s talk about WHY they\'re going to Spain specifically. Because your clients are '
    'going to ask you — and you need to have these answers ready.\n\n'
    'There are five core drivers. I call them the Five Forces of the Spanish Expat Boom."',
    speaker='GUIÓN EN INGLÉS:')

add_script_es(doc,
    '"Ahora hablemos de POR QUÉ van a España concretamente. Porque vuestros clientes os lo '
    'van a preguntar — y necesitáis tener estas respuestas preparadas.\n\n'
    'Hay cinco motores fundamentales. Yo los llamo las Cinco Fuerzas del Boom Expat en España."',
    speaker='GUIÓN EN ESPAÑOL:')

forces_en = [
    ('🟡 Force 1 — Cost of living (EN):',
     '"First: cost of living. The average North American family moving from a major metropolitan '
     'area saves between 40% and 60% on their monthly expenses. We\'re talking about groceries, '
     'utilities, transportation, dining out — everything. A €350,000 beachfront apartment in '
     'Alicante would cost over a million in Miami. That\'s not an estimate, that\'s a listing '
     'comparison I can show you right now."'),
    ('🟡 Force 2 — Healthcare (EN):',
     '"Second: healthcare. Spain ranks number 7 in the world for healthcare quality [WHO], '
     'compared to the US at number 37. A private comprehensive health plan for a family of four '
     'in Spain costs approximately €250 to €400 per month. In the United States, that same '
     'coverage is $2,000 or more. For your clients who are pre-retirement or dealing with high '
     'insurance premiums, this alone justifies the move."'),
    ('🟡 Force 3 — Safety (EN):',
     '"Third: safety. Spain is consistently ranked among the top 10 safest countries in the '
     'world [Global Peace Index]. Gun violence, as a risk factor, is virtually non-existent. '
     'For families with children, for retirees, for anyone concerned about personal security '
     '— Spain is a haven."'),
    ('🟡 Force 4 — Climate (EN):',
     '"Fourth: climate. The Costa Blanca — which is our primary market — has over 320 days of '
     'sunshine per year. Average winter temperatures are 14°C to 18°C — which is 57 to 65°F. '
     'For our Canadian clients especially, this is transformational. They\'re trading six months '
     'of darkness and frozen pipes for a Mediterranean lifestyle. No contest."'),
    ('🟡 Force 5 — Visas & residency (EN):',
     '"And fifth — and this is what\'s really driving the current wave — Spain\'s visa landscape '
     'has never been more favorable for North Americans.\n\n'
     'The Non-Lucrative Visa is perfect for retirees and those with passive income. The Digital '
     'Nomad Visa, launched in 2023, lets remote workers live in Spain legally for up to 5 years. '
     'The Golden Visa allows property investors of €500,000 or more to obtain full family residency.\n\n'
     'Our legal partner — Fuster & Associates, whom you\'ll meet shortly — handles all of this. '
     'Your clients don\'t need to figure it out alone. They just need to call us."'),
]
forces_es = [
    ('🟡 Fuerza 1 — Coste de vida (ES):',
     '"Primero: el coste de vida. La familia norteamericana media que se traslada desde una '
     'gran área metropolitana ahorra entre el 40% y el 60% en sus gastos mensuales. Hablamos '
     'de alimentación, suministros, transporte, comer fuera — absolutamente todo. Un apartamento '
     'en primera línea de playa en Alicante por 350.000 euros costaría más de un millón en Miami. '
     'No es una estimación, es una comparativa de listados que os puedo mostrar ahora mismo."'),
    ('🟡 Fuerza 2 — Sanidad (ES):',
     '"Segundo: la sanidad. España ocupa el puesto número 7 del mundo en calidad sanitaria [OMS], '
     'frente al puesto 37 de Estados Unidos. Un seguro médico privado completo para una familia '
     'de cuatro personas en España cuesta aproximadamente entre 250 y 400 euros al mes. En '
     'Estados Unidos, esa misma cobertura son 2.000 dólares o más. Para vuestros clientes que '
     'están en la previa de la jubilación o que pagan primas de seguro desorbitadas, esto '
     'por sí solo justifica el traslado."'),
    ('🟡 Fuerza 3 — Seguridad (ES):',
     '"Tercero: la seguridad. España se posiciona sistemáticamente entre los 10 países más '
     'seguros del mundo [Índice de Paz Global]. La violencia con armas de fuego, como factor '
     'de riesgo, es prácticamente inexistente. Para familias con hijos, para jubilados, para '
     'cualquiera que se preocupe por la seguridad personal — España es un refugio."'),
    ('🟡 Fuerza 4 — Clima (ES):',
     '"Cuarto: el clima. La Costa Blanca — nuestro mercado principal — tiene más de 320 días '
     'de sol al año. Las temperaturas medias en invierno son de 14°C a 18°C — entre 57 y 65°F. '
     'Para nuestros clientes canadienses especialmente, esto es transformador. Están cambiando '
     'seis meses de oscuridad y tuberías congeladas por un estilo de vida mediterráneo. '
     'No hay comparación posible."'),
    ('🟡 Fuerza 5 — Visados y residencia (ES):',
     '"Y quinto — y esto es lo que realmente está impulsando la ola actual — el panorama de '
     'visados en España nunca ha sido tan favorable para los norteamericanos.\n\n'
     'La Visa de No Lucrativa es perfecta para jubilados y personas con ingresos pasivos. '
     'La Visa de Nómada Digital, lanzada en 2023, permite a los trabajadores remotos vivir '
     'legalmente en España hasta 5 años. La Golden Visa permite a inversores inmobiliarios '
     'de 500.000 euros o más obtener la residencia completa para toda la familia.\n\n'
     'Nuestro partner legal — Fuster & Associates, a quien conoceréis enseguida — gestiona '
     'todo esto. Vuestros clientes no necesitan resolverlo solos. Solo necesitan llamarnos."'),
]
for (label_en, text_en), (label_es, text_es) in zip(forces_en, forces_es):
    add_paragraph(doc, label_en, bold=True, color=RGBColor(0xC8, 0x86, 0x00), size=11, space_before=8)
    add_script(doc, text_en)
    add_paragraph(doc, label_es, bold=True, color=RGBColor(0x22, 0x44, 0x88), size=11, space_before=4)
    add_script_es(doc, text_es)

# ── SLIDE 7 ──────────────────────────────────────────────────────────────────
add_heading2(doc, '🎬 SLIDE 7: "Your Clients Are Already Looking" — El Mercado que Ya Existe')
add_instruction(doc,
    '[Pantalla dividida: búsquedas de Google + perfiles de compradores tipo]')

add_script(doc,
    '"Here\'s the thing I want you to really internalize: your clients are NOT waiting for your '
    'permission to explore this. They are already googling \'how to move to Spain,\' they\'re '
    'already in Facebook groups, they\'re already watching YouTube videos of expats in Valencia '
    'and Alicante.\n\n'
    'The question is not IF they will explore it. The question is: will you be the person who '
    'guides them, or will they find someone else online and disappear from your pipeline forever?\n\n'
    'We are giving you the tools, the team, and the infrastructure to be THAT person.\n\n'
    'Let\'s talk about how."',
    speaker='GUIÓN EN INGLÉS:')

add_script_es(doc,
    '"Esto es lo que quiero que interioricéis de verdad: vuestros clientes NO están esperando '
    'vuestro permiso para explorar esto. Ya están buscando en Google "cómo mudarse a España", '
    'ya están en grupos de Facebook, ya están viendo vídeos de YouTube de expats en Valencia '
    'y Alicante.\n\n'
    'La pregunta no es SI van a explorarlo. La pregunta es: ¿vais a ser vosotros quienes '
    'les guíen, o van a encontrar a otra persona en internet y desaparecer de vuestro pipeline '
    'para siempre?\n\n'
    'Nosotros os damos las herramientas, el equipo y la infraestructura para ser ESA persona.\n\n'
    'Hablemos de cómo."',
    speaker='GUIÓN EN ESPAÑOL:')

add_separator(doc)

# ══════════════════════════════════════════════════════════════════════════════
# BLOQUE 3 — EL ECOSISTEMA 360°
# ══════════════════════════════════════════════════════════════════════════════
add_heading1(doc, 'BLOQUE 3 — EL ECOSISTEMA 360°')
add_paragraph(doc, '⏰ 20:00 – 32:00 | Hablan: Partners Invitados',
              italic=True, color=GREY, size=10)

# ── SLIDE 8 ──────────────────────────────────────────────────────────────────
add_heading2(doc, '🎬 SLIDE 8: INTRODUCCIÓN AL ECOSISTEMA')

add_script(doc,
    '"Now, the number one fear of a North American buyer going into a foreign market is: \'What '
    'if something goes wrong? What if I\'m scammed? What if I pay taxes I don\'t know about? '
    'What if the property title isn\'t clean?\'\n\n'
    'This fear is completely legitimate. And it\'s exactly why we built what we call our 360° '
    'Ecosystem.\n\n'
    'We don\'t sell properties in isolation. We deliver a complete, end-to-end service covering '
    'legal, financial, and property management support — all under one coordinated team.\n\n'
    'I\'d like to introduce you to three specialists who make this possible. Each of them is '
    'going to take about two minutes to explain what they do and why it matters for your clients. '
    'Pay attention — because this ecosystem is what makes our offer completely unique in the market.\n\n'
    'Let me start with perhaps the most critical piece: legal security."',
    speaker='GUIÓN EN INGLÉS (Agente 1 — introduce el bloque):')

add_script_es(doc,
    '"Ahora bien, el miedo número uno de un comprador norteamericano que se adentra en un '
    'mercado extranjero es: "¿Y si algo va mal? ¿Y si me estafan? ¿Y si pago impuestos que '
    'no conozco? ¿Y si el título de propiedad no está limpio?"\n\n'
    'Este miedo es completamente legítimo. Y es exactamente por eso que construimos lo que '
    'llamamos nuestro Ecosistema 360°.\n\n'
    'No vendemos propiedades de forma aislada. Ofrecemos un servicio completo, de principio '
    'a fin, que cubre el apoyo legal, financiero y de gestión de la propiedad — todo bajo '
    'un equipo coordinado.\n\n'
    'Me gustaría presentaros a tres especialistas que hacen esto posible. Cada uno va a '
    'tomarse unos dos minutos para explicar qué hace y por qué es importante para vuestros '
    'clientes. Prestad atención — porque este ecosistema es lo que hace que nuestra oferta '
    'sea completamente única en el mercado.\n\n'
    'Empecemos con quizás la pieza más crítica: la seguridad jurídica."',
    speaker='GUIÓN EN ESPAÑOL (Agente 1 — introduce el bloque):')

# ── SLIDE 9 ──────────────────────────────────────────────────────────────────
add_heading2(doc, '🎬 SLIDE 9: PARTNER LEGAL — Fuster & Associates')
add_instruction(doc, '[Logo de Fuster & Associates · 2 minutos]')

add_script(doc,
    '"Hello everyone, I\'m [Nombre] from Fuster & Associates. We are a bilingual law firm '
    'based in Alicante, and for the past [X] years we have exclusively specialized in assisting '
    'international buyers — particularly from the United States, Canada, and the United Kingdom '
    '— in legally securing their property purchase and their residency in Spain.\n\n'
    'When your client decides to buy in Spain, here\'s what our firm handles:\n\n'
    'Number one — the NIE number, which is the Spanish tax ID that every non-resident buyer '
    'needs. We have it done in days, not weeks.\n\n'
    'Number two — due diligence on every property: title deeds, outstanding debts, urban planning '
    'compliance. Your client will never buy a property with hidden surprises.\n\n'
    'Number three — the full drafting and review of purchase contracts, negotiated in English.\n\n'
    'Number four — residency applications. Whether it\'s a non-lucrative visa, a digital nomad '
    'visa, or a Golden Visa, our immigration team has a 98% success rate.\n\n'
    'Our job is to make the legal process completely invisible to your client — they experience '
    'it as smooth, safe, and professional. You, as the referring Realtor, look like a hero '
    'because you sent them to the right team.\n\n'
    'We look forward to protecting your clients\' most important investment."',
    speaker='GUIÓN EN INGLÉS (Representante de Fuster):')

add_script_es(doc,
    '"Hola a todos, soy [Nombre] de Fuster & Associates. Somos un despacho de abogados '
    'bilingüe con sede en Alicante, y durante los últimos [X] años nos hemos especializado '
    'exclusivamente en asistir a compradores internacionales — especialmente de Estados Unidos, '
    'Canadá y Reino Unido — en la adquisición legal de su propiedad y en la obtención de '
    'su residencia en España.\n\n'
    'Cuando vuestro cliente decide comprar en España, esto es lo que gestiona nuestro despacho:\n\n'
    'Número uno — el NIE, que es el número de identificación fiscal español que necesita '
    'todo comprador no residente. Nosotros lo tramitamos en días, no en semanas.\n\n'
    'Número dos — la diligencia debida sobre cada propiedad: escrituras, deudas pendientes, '
    'cumplimiento urbanístico. Vuestro cliente nunca comprará una propiedad con sorpresas ocultas.\n\n'
    'Número tres — la redacción y revisión completa de los contratos de compraventa, '
    'negociados en inglés.\n\n'
    'Número cuatro — las solicitudes de residencia. Ya sea la visa de no lucrativa, la visa '
    'de nómada digital o la Golden Visa, nuestro equipo de inmigración tiene un 98% de '
    'tasa de éxito.\n\n'
    'Nuestro trabajo es hacer que el proceso legal sea completamente invisible para vuestro '
    'cliente — que lo viva como algo fluido, seguro y profesional. Vosotros, como Realtors '
    'referentes, quedáis como héroes por haberles enviado al equipo correcto.\n\n'
    'Esperamos con entusiasmo proteger la inversión más importante de vuestros clientes."',
    speaker='GUIÓN EN ESPAÑOL (Representante de Fuster):')

add_instruction(doc, '[Agente 1 regresa:]')
add_script(doc,
    '"Thank you. Fuster is the reason our buyers sleep soundly at night. Now let\'s talk about '
    'something else that keeps buyers awake: how to pay for the property."')
add_script_es(doc,
    '"Gracias. Fuster es la razón por la que nuestros compradores duermen tranquilos. '
    'Ahora hablemos de algo más que también les quita el sueño: cómo financiar la propiedad."')

# ── SLIDE 10 ──────────────────────────────────────────────────────────────────
add_heading2(doc, '🎬 SLIDE 10: PARTNER FINANCIERO — UCI (Unión de Créditos Inmobiliarios)')
add_instruction(doc, '[Logo de UCI · 2 minutos]')

add_script(doc,
    '"Thank you. I\'m [Nombre] from UCI — Unión de Créditos Inmobiliarios. We are one of '
    'Spain\'s leading mortgage specialists, and we have a dedicated product line for '
    'non-resident buyers, including Americans and Canadians.\n\n'
    'There\'s a common misconception that foreigners can\'t get a mortgage in Spain. That is '
    'simply not true.\n\n'
    'Non-resident buyers from the US and Canada can access mortgages for up to 70% of the '
    'property value at very competitive fixed rates — currently in the range of 2.8% to 3.5% '
    '— which are substantially lower than current US mortgage rates.\n\n'
    'Our process is designed for international clients: we conduct income verification based on '
    'US or Canadian tax returns, we have English-speaking advisors, and we can issue a mortgage '
    'pre-approval in as little as 48 to 72 hours.\n\n'
    'Why does this matter for you as a Realtor? Because many of your clients who say \'I\'d '
    'love to buy in Spain but I can\'t afford to pay cash\' actually CAN buy in Spain with '
    'financing. We remove the cash-only barrier.\n\n'
    'We are here to make sure price is never the reason your client says no."',
    speaker='GUIÓN EN INGLÉS (Representante de UCI):')

add_script_es(doc,
    '"Gracias. Soy [Nombre] de UCI — Unión de Créditos Inmobiliarios. Somos uno de los '
    'principales especialistas hipotecarios de España, y disponemos de una línea de '
    'productos dedicada específicamente a compradores no residentes, incluidos americanos '
    'y canadienses.\n\n'
    'Existe un malentendido común: que los extranjeros no pueden obtener hipoteca en España. '
    'Eso es sencillamente falso.\n\n'
    'Los compradores no residentes de EE.UU. y Canadá pueden acceder a hipotecas de hasta '
    'el 70% del valor de la propiedad a tipos fijos muy competitivos — actualmente en el '
    'rango del 2,8% al 3,5% — que son sustancialmente más bajos que los tipos hipotecarios '
    'actuales en EE.UU.\n\n'
    'Nuestro proceso está diseñado para clientes internacionales: realizamos la verificación '
    'de ingresos basándonos en las declaraciones fiscales americanas o canadienses, contamos '
    'con asesores de habla inglesa y podemos emitir una preaprobación hipotecaria en tan '
    'solo 48 a 72 horas.\n\n'
    '¿Por qué esto importa para vosotros como Realtors? Porque muchos de vuestros clientes '
    'que dicen "me encantaría comprar en España pero no puedo permitirme pagar al contado" '
    'en realidad SÍ pueden comprar en España con financiación. Nosotros eliminamos la barrera '
    'del pago en efectivo.\n\n'
    'Estamos aquí para asegurarnos de que el precio nunca sea la razón por la que vuestro '
    'cliente diga que no."',
    speaker='GUIÓN EN ESPAÑOL (Representante de UCI):')

add_instruction(doc, '[Agente 2 regresa:]')
add_script(doc,
    '"Remarkable, right? European mortgage rates that are significantly lower than what your '
    'clients are seeing at home. Now, here\'s a question I get a lot: \'What if my client wants '
    'to buy the property but only live there part of the year? Can they generate income from it?\' '
    'Absolutely — and here\'s how."')
add_script_es(doc,
    '"Sorprendente, ¿verdad? Tipos hipotecarios europeos significativamente más bajos que '
    'los que ven vuestros clientes en casa. Ahora bien, aquí hay una pregunta que recibo '
    'constantemente: "¿Y si mi cliente quiere comprar la propiedad pero solo vivir allí '
    'parte del año? ¿Puede generar ingresos con ella?" Absolutamente — y os explico cómo."')

# ── SLIDE 11 ──────────────────────────────────────────────────────────────────
add_heading2(doc, '🎬 SLIDE 11: PARTNER DE GESTIÓN — Inmomás Holidays')
add_instruction(doc, '[Logo de Inmomás Holidays · 2 minutos]')

add_script(doc,
    '"Hi everyone, I\'m [Nombre] from Inmomás Holidays — we are the property management and '
    'vacation rental arm of the RE/MAX Inmomás group.\n\n'
    'Many North American buyers are buying in Spain not just as a primary residence, but as an '
    'investment — they want to use the property 4 to 6 weeks a year and have it generating '
    'income the rest of the time. That\'s exactly what we manage.\n\n'
    'Here\'s what we do: we handle the tourist rental license — which is required by Spanish law '
    'and quite complex to obtain. We list the property on all major platforms — Airbnb, '
    'Booking.com, VRBO. We manage check-in, cleaning, maintenance, and guest communication. '
    'Our clients receive a monthly report and a bank transfer, and they don\'t lift a finger.\n\n'
    'What are the yields? On a well-located Costa Blanca property, we consistently deliver net '
    'rental yields of between 5% and 8% per year after our management fee.\n\n'
    'So your client\'s property doesn\'t just appreciate — it pays for itself while they\'re '
    'not there. That\'s a very compelling story to tell your investor clients."',
    speaker='GUIÓN EN INGLÉS (Representante de Inmomás Holidays):')

add_script_es(doc,
    '"Hola a todos, soy [Nombre] de Inmomás Holidays — somos el brazo de gestión de '
    'propiedades y alquiler vacacional del grupo RE/MAX Inmomás.\n\n'
    'Muchos compradores norteamericanos compran en España no solo como residencia principal, '
    'sino como inversión — quieren usar la propiedad 4 o 6 semanas al año y que genere '
    'ingresos el resto del tiempo. Eso es exactamente lo que nosotros gestionamos.\n\n'
    'Esto es lo que hacemos: tramitamos la licencia de alquiler turístico — que exige la '
    'ley española y es bastante compleja de obtener. Publicamos la propiedad en todas las '
    'plataformas principales — Airbnb, Booking.com, VRBO. Gestionamos el check-in, la '
    'limpieza, el mantenimiento y la comunicación con los huéspedes. Nuestros clientes '
    'reciben un informe mensual y una transferencia bancaria, sin mover un dedo.\n\n'
    '¿Cuáles son las rentabilidades? En una propiedad bien ubicada en la Costa Blanca, '
    'entregamos consistentemente rentabilidades netas de alquiler de entre el 5% y el 8% '
    'anual tras nuestra comisión de gestión.\n\n'
    'Así que la propiedad de vuestro cliente no solo se revaloriza — se paga sola mientras '
    'ellos no están. Eso es una historia muy convincente para contarles a vuestros '
    'clientes inversores."',
    speaker='GUIÓN EN ESPAÑOL (Representante de Inmomás Holidays):')

add_instruction(doc, '[Agente 1 regresa:]')
add_script(doc,
    '"Excellent. Now that you\'ve met our three key partners, let me tell you about the '
    'experience that — in our data — closes more deals than any brochure or website ever could."')
add_script_es(doc,
    '"Excelente. Ahora que ya conocéis a nuestros tres partners clave, dejadme hablaros '
    'de la experiencia que — según nuestros datos — cierra más operaciones que cualquier '
    'folleto o página web jamás podría conseguir."')

# ── SLIDE 12 ──────────────────────────────────────────────────────────────────
add_heading2(doc, '🎬 SLIDE 12: EL VIP PROPERTY TOUR — La Experiencia que Cierra la Venta')
add_instruction(doc,
    '[Fotos: clientes en tour de propiedades, cenas en restaurantes, vista al Mediterráneo]')

add_script(doc,
    '"I want to close this section with something that I think is the most powerful sales tool '
    'in our entire program — and it costs your client nothing extra.\n\n'
    'It\'s called the VIP Property Tour.\n\n'
    'When a client is seriously interested, we invite them to Spain for a 4 to 5 day fully '
    'organized visit. We pick them up from the airport. We take them to carefully selected '
    'properties. We arrange meetings with the legal team and the mortgage specialist. We organize '
    'dinners where they meet other American and Canadian expats who already live here — people '
    'they can ask \'is this really as good as it sounds?\' And the answer, every single time, '
    'is: \'Yes. And I wish I\'d done it sooner.\'\n\n'
    'The close rate on clients who come on a VIP tour is over 70%. They arrive as prospects. '
    'They leave as buyers.\n\n'
    'You make the introduction. We do the rest. You earn your 25%."',
    speaker='GUIÓN EN INGLÉS (Agente 2):')

add_script_es(doc,
    '"Quiero cerrar esta sección con algo que creo que es la herramienta de ventas más '
    'poderosa de todo nuestro programa — y no le cuesta nada adicional a vuestro cliente.\n\n'
    'Se llama el VIP Property Tour.\n\n'
    'Cuando un cliente tiene un interés serio, le invitamos a España para una visita de '
    '4 a 5 días completamente organizada. Le recogemos en el aeropuerto. Le llevamos a '
    'propiedades cuidadosamente seleccionadas. Organizamos reuniones con el equipo legal '
    'y el especialista hipotecario. Organizamos cenas donde conoce a otros expats americanos '
    'y canadienses que ya viven aquí — personas a quienes puede preguntarles "¿de verdad '
    'es tan bueno como suena?" Y la respuesta, absolutamente siempre, es: "Sí. '
    'Y ojalá lo hubiera hecho antes."\n\n'
    'La tasa de cierre de los clientes que vienen al VIP Tour supera el 70%. Llegan como '
    'interesados. Se van como compradores.\n\n'
    'Vosotros hacéis la presentación. Nosotros hacemos el resto. Vosotros cobráis vuestro 25%."',
    speaker='GUIÓN EN ESPAÑOL (Agente 2):')

add_separator(doc)

# ══════════════════════════════════════════════════════════════════════════════
# BLOQUE 4 — LA FÓRMULA FINANCIERA
# ══════════════════════════════════════════════════════════════════════════════
add_heading1(doc, 'BLOQUE 4 — LA FÓRMULA FINANCIERA')
add_paragraph(doc, '⏰ 32:00 – 40:00 | Presentan: Agentes de RE/MAX Inmomás',
              italic=True, color=GREY, size=10)

# ── SLIDE 13 ──────────────────────────────────────────────────────────────────
add_heading2(doc, '🎬 SLIDE 13: "The Commission Blueprint" — Tu 25%')
add_instruction(doc,
    '[Infografía clara con el flujo de comisión: Buyer → 5% Fee → 25% a Realtor Referente]')

add_script(doc,
    '"Alright. This is the part everyone\'s been waiting for. Let\'s talk money.\n\n'
    'Here\'s how the commission structure works — and I\'m going to be 100% transparent '
    'because I know that\'s what you need as a professional.\n\n'
    'In Spain, our team manages the full buyer-side service. We charge the buyer a comprehensive '
    'service fee of 5% of the property purchase price. This fee covers everything: property '
    'search, legal coordination, financing assistance, negotiation, closing, and post-sale support.\n\n'
    'As the referring Realtor — meaning you introduced us to your client — you receive 25% of '
    'that total fee.\n\n'
    'Let me do the math for you."',
    speaker='GUIÓN EN INGLÉS (Agente 1):')

add_script_es(doc,
    '"Muy bien. Esta es la parte que todo el mundo estaba esperando. Hablemos de dinero.\n\n'
    'Así funciona la estructura de comisiones — y voy a ser 100% transparente porque sé '
    'que eso es lo que necesitáis como profesionales.\n\n'
    'En España, nuestro equipo gestiona el servicio completo del lado del comprador. '
    'Cobramos al comprador una tarifa de servicio integral del 5% del precio de compra '
    'de la propiedad. Esta tarifa lo cubre todo: búsqueda de propiedad, coordinación legal, '
    'asistencia con la financiación, negociación, cierre y soporte postventa.\n\n'
    'Como Realtor referente — es decir, vosotros que nos habéis presentado a vuestro '
    'cliente — recibís el 25% de esa tarifa total.\n\n'
    'Dejadme hacer los cálculos por vosotros."',
    speaker='GUIÓN EN ESPAÑOL (Agente 1):')

# ── SLIDE 14 ──────────────────────────────────────────────────────────────────
add_heading2(doc, '🎬 SLIDE 14: TABLA DE GANANCIAS REALES — "Real Earnings"')
add_instruction(doc, '[Tabla grande y visual con tres ejemplos de propiedades]')

add_script(doc,
    '"Let\'s look at three real scenarios:"',
    speaker='GUIÓN EN INGLÉS:')
add_script_es(doc,
    '"Veamos tres escenarios reales:"',
    speaker='GUIÓN EN ESPAÑOL:')

# Tabla de ganancias
earnings_data = [
    ('Apartamento en playa (Alicante)', '€350,000', '~$380,000 USD', '€17,500', '€4,375'),
    ('Villa premium con piscina (Golden Visa)', '€600,000', '~$652,000 USD', '€30,000', '€7,500'),
    ('Villa de lujo en la costa', '€1,200,000', '~$1,304,000 USD', '€60,000', '€15,000'),
]

t3 = doc.add_table(rows=1 + len(earnings_data), cols=5)
t3.alignment = WD_TABLE_ALIGNMENT.LEFT
t3.style = 'Table Grid'
set_table_borders(t3)

earn_hdrs = ('Propiedad', 'Precio (€)', 'Equivalente USD', 'Fee 5%', 'Tu 25% ✓')
for ci, h in enumerate(earn_hdrs):
    c = t3.rows[0].cells[ci]
    set_cell_bg(c, 'C8102E')
    run = c.paragraphs[0].add_run(h)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for i, row_data in enumerate(earnings_data, start=1):
    row = t3.rows[i].cells
    bg = 'FFF0F0' if i % 2 == 0 else 'FFFFFF'
    for ci, txt in enumerate(row_data):
        set_cell_bg(row[ci], bg)
        run = row[ci].paragraphs[0].add_run(txt)
        run.font.size = Pt(10)
        row[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if ci == 4:
            run.bold = True
            run.font.color.rgb = RED
        elif ci == 0:
            run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
            row[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            run.font.color.rgb = BLACK

t3.columns[0].width = Cm(5.5)
t3.columns[1].width = Cm(2.5)
t3.columns[2].width = Cm(3.0)
t3.columns[3].width = Cm(2.0)
t3.columns[4].width = Cm(2.0)

doc.add_paragraph()

add_script(doc,
    '"Now think about this: how many hours do you currently spend to earn $7,500 in your local '
    'market? Showings, open houses, negotiations, inspections, contingencies, appraisals...\n\n'
    'With us, you spend five minutes filling out our partner portal. We handle the rest. And '
    'you get paid."')
add_script_es(doc,
    '"Pensad en esto: ¿cuántas horas dedicáis actualmente para ganar 7.500 dólares en vuestro '
    'mercado local? Visitas, open houses, negociaciones, inspecciones, contingencias, tasaciones...\n\n'
    'Con nosotros, invertís cinco minutos rellenando el formulario en nuestro portal de '
    'partners. Nosotros nos encargamos del resto. Y vosotros cobráis."')

add_script(doc,
    '"I\'m not suggesting you abandon your local market. Absolutely not. We\'re not asking you '
    'to change anything about your business. What we ARE offering is a new, parallel income '
    'stream that activates every time one of your clients expresses interest in Spain — which, '
    'based on the chat messages you\'re sending right now, happens more often than you might think.\n\n'
    'This isn\'t a side hustle. This is a referral network upgrade."',
    speaker='GUIÓN EN INGLÉS (continuación — tono cómplice):')
add_script_es(doc,
    '"No os estoy sugiriendo que abandonéis vuestro mercado local. En absoluto. No os pedimos '
    'que cambiéis nada de vuestro negocio. Lo que SÍ os ofrecemos es una nueva fuente de '
    'ingresos paralela que se activa cada vez que uno de vuestros clientes expresa interés '
    'en España — lo cual, a juzgar por los mensajes que estáis enviando en el chat ahora '
    'mismo, ocurre con más frecuencia de la que creéis.\n\n'
    'Esto no es un ingreso extra puntual. Es una mejora de vuestra red de referencias."',
    speaker='GUIÓN EN ESPAÑOL (continuación — tono cómplice):')

# ── SLIDE 15 ──────────────────────────────────────────────────────────────────
add_heading2(doc, '🎬 SLIDE 15: EL PROCESO — "How Simple Is It?"')
add_instruction(doc,
    '[Infografía de 4 pasos: 1. Introduce → 2. Register → 3. We Handle → 4. You Get Paid]')

add_script(doc,
    '"Here\'s the entire process in four steps:\n\n'
    'Step one: your client mentions Spain. You say: \'I actually have the perfect team for that. '
    'Let me connect you.\'\n\n'
    'Step two: you log into our partner portal — which takes two minutes to register — and you '
    'enter your client\'s basic information.\n\n'
    'Step three: we take over completely. Your client is contacted by our bilingual team within '
    '24 hours. We manage the search, the tours, the legal process, the mortgage, everything. '
    'You can track every step in real time on your dashboard — like a Kanban board for your referrals.\n\n'
    'Step four: the deal closes. You receive your referral fee — in your currency of choice, via '
    'international wire transfer, with a formal referral agreement signed digitally.\n\n'
    'That\'s it. Four steps. Zero additional licensing. Zero extra hours.\n\n'
    'You just extended your business to Europe."',
    speaker='GUIÓN EN INGLÉS:')

add_script_es(doc,
    '"Este es todo el proceso en cuatro pasos:\n\n'
    'Paso uno: vuestro cliente menciona España. Vosotros decís: "Tengo el equipo perfecto '
    'para eso. Os pongo en contacto."\n\n'
    'Paso dos: entráis en nuestro portal de partners — el registro lleva dos minutos — '
    'e introducís la información básica de vuestro cliente.\n\n'
    'Paso tres: nosotros nos hacemos cargo por completo. Nuestro equipo bilingüe contacta '
    'al cliente en menos de 24 horas. Gestionamos la búsqueda, las visitas, el proceso legal, '
    'la hipoteca, todo. Podéis seguir cada paso en tiempo real en vuestro panel de control '
    '— como un tablero Kanban para vuestras referencias.\n\n'
    'Paso cuatro: la operación cierra. Recibís vuestra comisión de referencia — en vuestra '
    'divisa preferida, mediante transferencia internacional, con un acuerdo formal de '
    'referencia firmado digitalmente.\n\n'
    'Eso es todo. Cuatro pasos. Sin licencias adicionales. Sin horas extra.\n\n'
    'Acabáis de extender vuestro negocio a Europa."',
    speaker='GUIÓN EN ESPAÑOL:')

add_separator(doc)

# ══════════════════════════════════════════════════════════════════════════════
# BLOQUE 5 — DEMO DEL PORTAL Y CIERRE
# ══════════════════════════════════════════════════════════════════════════════
add_heading1(doc, 'BLOQUE 5 — DEMO DEL PORTAL Y CIERRE')
add_paragraph(doc, '⏰ 40:00 – 52:00 | Presentan: Agentes de RE/MAX Inmomás',
              italic=True, color=GREY, size=10)

# ── SLIDE 16 ──────────────────────────────────────────────────────────────────
add_heading2(doc, '🎬 SLIDE 16: DEMO EN VIVO DEL PORTAL PARTNER')
add_instruction(doc,
    '[Compartir pantalla: demostración del portal web inmomas-international.com/app.html]')

add_script(doc,
    '"Let me show you exactly what you\'ll be working with. I\'m going to share my screen now.\n\n'
    'This is our RE/MAX Inmomás International partner portal. [Mostrar pantalla de login / dashboard]\n\n'
    'Here\'s what you see when you log in as a partner Realtor: your personal dashboard with all '
    'your referred clients, their current status in the buying process — whether they\'re in the '
    'initial consultation phase, the property search phase, the offer phase, or post-signing — '
    'and a direct communication channel with your assigned agent here in Spain.\n\n'
    'You don\'t need to send emails asking \'hey, what happened with my client?\' You can see it. '
    'In real time. Right here.\n\n'
    'On the right side, you\'ll see your earnings tracker — every referral, every pending '
    'commission, every payment processed.\n\n'
    '[Mostrar sección de documentos]\n\n'
    'Here\'s your co-branded marketing toolkit: your personalized landing page link, your '
    'bilingual email templates, your social media graphics — ready to download and share in minutes.\n\n'
    '[Mostrar sección de acuerdo de referidos]\n\n'
    'And here\'s the referral agreement — pre-signed by RE/MAX Inmomás, with space for your '
    'digital signature. It\'s a legal document that protects you, protects your client, and '
    'defines your fee. You\'ll have it in your inbox within 24 hours of submitting your first referral.\n\n'
    'This is a professional system built for professional Realtors."',
    speaker='GUIÓN EN INGLÉS (Agente 2):')

add_script_es(doc,
    '"Dejadme mostraros exactamente con qué vais a trabajar. Voy a compartir mi pantalla ahora.\n\n'
    'Este es nuestro portal de partners de RE/MAX Inmomás International. [Mostrar pantalla de login / dashboard]\n\n'
    'Esto es lo que veis cuando entráis como Realtor partner: vuestro panel personal con '
    'todos vuestros clientes referidos, su estado actual en el proceso de compra — si están '
    'en la fase de consulta inicial, en la fase de búsqueda de propiedad, en la fase de '
    'oferta o ya post-firma — y un canal de comunicación directa con vuestro agente asignado '
    'aquí en España.\n\n'
    'No necesitáis enviar emails preguntando "¿qué ha pasado con mi cliente?". Podéis '
    'verlo vosotros mismos. En tiempo real. Aquí mismo.\n\n'
    'En el lado derecho, veréis vuestro rastreador de ganancias — cada referencia, cada '
    'comisión pendiente, cada pago procesado.\n\n'
    '[Mostrar sección de documentos]\n\n'
    'Aquí está vuestro kit de marketing co-branded: vuestro enlace de landing page '
    'personalizado, vuestras plantillas de email bilingüe, vuestros gráficos para '
    'redes sociales — listos para descargar y compartir en minutos.\n\n'
    '[Mostrar sección de acuerdo de referidos]\n\n'
    'Y aquí está el acuerdo de referencia — pre-firmado por RE/MAX Inmomás, con espacio '
    'para vuestra firma digital. Es un documento legal que os protege a vosotros, protege '
    'a vuestro cliente y define vuestra tarifa. Lo tendréis en vuestro correo en menos '
    'de 24 horas tras enviar vuestra primera referencia.\n\n'
    'Este es un sistema profesional construido para Realtors profesionales."',
    speaker='GUIÓN EN ESPAÑOL (Agente 2):')

# ── SLIDE 17 ──────────────────────────────────────────────────────────────────
add_heading2(doc, '🎬 SLIDE 17: CALL TO ACTION — "Your Next Step"')
add_instruction(doc,
    '[Diapositiva con QR code, URL del portal, y oferta especial para asistentes al webinario]')

add_script(doc,
    '"We\'re going to open up for questions in just a moment. But first, let me tell you what '
    'we\'d like you to do right now, before you close this tab and go back to your day.\n\n'
    'Step one: Scan the QR code on your screen — or visit [URL] — and register as an official '
    'partner. It takes less than two minutes. It costs you absolutely nothing.\n\n'
    'Step two: download your bilingual Marketing Kit. It includes your co-branded landing page, '
    'three ready-to-send email templates, and five social media posts you can publish this week.\n\n'
    'Step three: think of one client — just one — who has mentioned Spain. Send me their name '
    'in the chat right now. Not a commitment. Just a name. I\'ll personally follow up with you '
    'after the webinar.\n\n'
    'And that\'s it. That\'s all we need from you to get started.\n\n'
    '[Pausa emocional]\n\n'
    'I want to close with something personal. Before I moved to Spain, I looked at opportunities '
    'like this one and I thought: \'Sounds too good to be true.\' I understand that skepticism. '
    'I really do.\n\n'
    'But here\'s what I know now that I didn\'t know then: the North American expat market in '
    'Spain is not a niche. It\'s not a trend. It\'s a generational movement — millions of people '
    'redefining what quality of life means, and choosing Spain to live it.\n\n'
    'You have clients who are part of that movement. We have the platform, the team, and the '
    'track record to serve them at the highest level.\n\n'
    'Together, we win.\n\n'
    'Thank you so much for being here today. Let\'s open up for Q&A."',
    speaker='GUIÓN EN INGLÉS (Agente 1 — cierre poderoso):')

add_script_es(doc,
    '"Vamos a abrir el turno de preguntas en un momento. Pero primero, dejadme deciros '
    'qué nos gustaría que hicierais ahora mismo, antes de cerrar esta pestaña y volver '
    'a vuestro día.\n\n'
    'Paso uno: escaneád el código QR en vuestra pantalla — o visitad [URL] — y registraos '
    'como partners oficiales. Lleva menos de dos minutos. No os cuesta absolutamente nada.\n\n'
    'Paso dos: descargad vuestro Kit de Marketing bilingüe. Incluye vuestra landing page '
    'co-branded, tres plantillas de email listas para enviar y cinco publicaciones para '
    'redes sociales que podéis publicar esta misma semana.\n\n'
    'Paso tres: pensad en un cliente — solo uno — que haya mencionado España. Mandadme '
    'su nombre en el chat ahora mismo. Sin compromiso. Solo un nombre. Yo personalmente '
    'me pondré en contacto con vosotros después del webinario.\n\n'
    'Y eso es todo. Eso es todo lo que necesitamos de vosotros para empezar.\n\n'
    '[Pausa emocional]\n\n'
    'Quiero cerrar con algo personal. Antes de mudarme a España, miraba oportunidades '
    'como esta y pensaba: "Suena demasiado bien para ser verdad." Entiendo ese escepticismo. '
    'De verdad que lo entiendo.\n\n'
    'Pero esto es lo que sé ahora y que no sabía entonces: el mercado expat norteamericano '
    'en España no es un nicho. No es una tendencia. Es un movimiento generacional — millones '
    'de personas redefiniendo lo que significa calidad de vida, y eligiendo España para vivirla.\n\n'
    'Tenéis clientes que forman parte de ese movimiento. Nosotros tenemos la plataforma, '
    'el equipo y la trayectoria para servirles al más alto nivel.\n\n'
    'Juntos, ganamos.\n\n'
    'Muchísimas gracias por estar aquí hoy. Abrimos el turno de Q&A."',
    speaker='GUIÓN EN ESPAÑOL (Agente 1 — cierre poderoso):')

add_separator(doc)

# ══════════════════════════════════════════════════════════════════════════════
# Q&A — PREGUNTAS FRECUENTES
# ══════════════════════════════════════════════════════════════════════════════
add_heading1(doc, '💬 SESIÓN Q&A — PREGUNTAS FRECUENTES Y RESPUESTAS MODELO')
add_paragraph(doc, '⏰ 46:00 – 52:00', italic=True, color=GREY, size=10)

qas = [
    (
        '"How does my client pay you if they\'re buying in euros?"',
        '"Great question. All transactions are done in euros — that\'s standard for Spanish '
        'property. Most of our North American clients use currency exchange specialists like '
        'Wise or OFX to transfer funds at competitive rates. UCI, our mortgage partner, also '
        'supports US/Canadian taxpayers with euro-denominated financing. We provide full guidance '
        'on this — your client won\'t have to figure it out alone."',
        '"¿Cómo paga mi cliente si compra en euros?"',
        '"Excelente pregunta. Todas las transacciones se realizan en euros — eso es estándar '
        'en el mercado inmobiliario español. La mayoría de nuestros clientes norteamericanos '
        'utilizan especialistas en cambio de divisas como Wise o OFX para transferir fondos '
        'a tipos competitivos. UCI, nuestro partner hipotecario, también da soporte a '
        'contribuyentes de EE.UU. y Canadá con financiación denominada en euros. '
        'Les orientamos en todo el proceso — vuestro cliente no tendrá que resolverlo solo."'
    ),
    (
        '"Do I need to be licensed in Spain to receive a referral fee?"',
        '"No. In Spain, the real estate agency system does not require the same licensing '
        'structure as NAR in the US or CREA in Canada. The referral fee is paid to you as a '
        'business-to-business service fee under a referral agreement — not as a commission split. '
        'We recommend you verify with your own broker that this is permissible under your '
        'brokerage\'s policy, which in the vast majority of cases it is for international referrals. '
        'Our legal team can provide documentation to support this if needed."',
        '"¿Necesito licencia en España para recibir una comisión de referencia?"',
        '"No. En España, el sistema de agencias inmobiliarias no requiere la misma estructura '
        'de licencias que la NAR en EE.UU. o la CREA en Canadá. La comisión de referencia '
        'se os abona como una tarifa de servicio B2B bajo un acuerdo de referencia — no como '
        'un reparto de comisión. Os recomendamos verificar con vuestro broker que esto está '
        'permitido bajo la política de vuestra agencia, lo cual en la gran mayoría de casos '
        'es así para referencias internacionales. Nuestro equipo legal puede proporcionar '
        'documentación de soporte si la necesitáis."'
    ),
    (
        '"What if my client decides not to buy? Do I still get paid?"',
        '"The referral fee is paid upon successful closing. If the client decides not to buy, '
        'there is no fee — from us or from you. There are no upfront costs, no retainers, no '
        'risk to you whatsoever. You invest zero. If it closes, you earn. That\'s the deal."',
        '"¿Y si mi cliente decide no comprar? ¿Cobro igualmente?"',
        '"La comisión de referencia se paga al cierre exitoso de la operación. Si el cliente '
        'decide no comprar, no hay ningún cargo — ni de nuestra parte ni de la vuestra. '
        'No hay costes iniciales, ni honorarios de retención, ni ningún riesgo para vosotros. '
        'Vosotros invertís cero. Si se cierra, cobráis. Así de sencillo es el acuerdo."'
    ),
    (
        '"How do I know my client is protected legally?"',
        '"Our legal partner — Fuster & Associates — has been protecting international buyers in '
        'Spain for over [X] years. They conduct full due diligence on every property: title check, '
        'urban planning verification, outstanding debts, compliance with all local regulations. '
        'Every purchase goes through an independent notary. Your client is protected at every step, '
        'and you can communicate directly with the legal team in English."',
        '"¿Cómo sé que mi cliente está protegido legalmente?"',
        '"Nuestro partner legal — Fuster & Associates — lleva más de [X] años protegiendo '
        'a compradores internacionales en España. Realizan una diligencia debida completa '
        'sobre cada propiedad: verificación del título, comprobación urbanística, deudas '
        'pendientes, cumplimiento de toda la normativa local. Cada compraventa pasa por '
        'un notario independiente. Vuestro cliente está protegido en cada paso, y podéis '
        'comunicaros directamente con el equipo legal en inglés."'
    ),
    (
        '"What are the typical taxes a buyer has to pay in Spain?"',
        '"For resale properties, buyers pay a transfer tax — called ITP — which varies by region, '
        'typically between 8% and 10% of the purchase price. For new builds, it\'s VAT at 10% '
        'plus 1.5% stamp duty. These taxes are factored into every budget we provide to clients '
        'from day one — no surprises. Fuster\'s team provides a complete cost breakdown before '
        'any offer is made."',
        '"¿Qué impuestos paga habitualmente un comprador en España?"',
        '"Para propiedades de segunda mano, los compradores pagan un impuesto de transmisiones '
        'patrimoniales — llamado ITP — que varía según la región, generalmente entre el 8% '
        'y el 10% del precio de compra. Para obra nueva, es el IVA al 10% más el 1,5% de '
        'actos jurídicos documentados. Estos impuestos están incluidos en cada presupuesto '
        'que facilitamos a los clientes desde el primer día — sin sorpresas. El equipo de '
        'Fuster proporciona un desglose completo de costes antes de realizar cualquier oferta."'
    ),
    (
        '"Can Canadian clients also get the Golden Visa?"',
        '"Absolutely. The Golden Visa program is available to non-EU citizens regardless of '
        'nationality — Americans, Canadians, and others. A minimum real estate investment of '
        '€500,000 unlocks full family residency, including for spouse and dependent children. '
        'Processing time is currently 3 to 6 months. Fuster & Associates handles the entire '
        'application."',
        '"¿Los clientes canadienses también pueden obtener la Golden Visa?"',
        '"Por supuesto. El programa de Golden Visa está disponible para ciudadanos no '
        'comunitarios independientemente de su nacionalidad — americanos, canadienses y otros. '
        'Una inversión inmobiliaria mínima de 500.000 euros abre la puerta a la residencia '
        'completa para toda la familia, incluyendo cónyuge e hijos dependientes. El tiempo '
        'de tramitación actual es de 3 a 6 meses. Fuster & Associates gestiona íntegramente '
        'la solicitud."'
    ),
]

for i, (q_en, a_en, q_es, a_es) in enumerate(qas, start=1):
    # EN
    p_q = doc.add_paragraph()
    p_q.paragraph_format.space_before = Pt(10)
    p_q.paragraph_format.space_after  = Pt(3)
    run_num  = p_q.add_run(f'Q{i} (EN):  ')
    run_num.bold = True
    run_num.font.size      = Pt(11)
    run_num.font.color.rgb = RED
    run_preg = p_q.add_run(q_en)
    run_preg.italic = True
    run_preg.font.size      = Pt(11)
    run_preg.font.color.rgb = RGBColor(0x22, 0x22, 0x66)

    p_r = doc.add_paragraph()
    p_r.paragraph_format.left_indent  = Cm(1.0)
    p_r.paragraph_format.space_before = Pt(2)
    p_r.paragraph_format.space_after  = Pt(3)
    run_label = p_r.add_run('Answer: ')
    run_label.bold = True
    run_label.font.size      = Pt(10.5)
    run_label.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    run_resp = p_r.add_run(a_en)
    run_resp.italic = True
    run_resp.font.size      = Pt(10.5)
    run_resp.font.color.rgb = BLACK

    # ES
    p_q2 = doc.add_paragraph()
    p_q2.paragraph_format.space_before = Pt(4)
    p_q2.paragraph_format.space_after  = Pt(3)
    run_num2  = p_q2.add_run(f'P{i} (ES):  ')
    run_num2.bold = True
    run_num2.font.size      = Pt(11)
    run_num2.font.color.rgb = BLUE
    run_preg2 = p_q2.add_run(q_es)
    run_preg2.italic = True
    run_preg2.font.size      = Pt(11)
    run_preg2.font.color.rgb = RGBColor(0x11, 0x11, 0x66)

    p_r2 = doc.add_paragraph()
    p_r2.paragraph_format.left_indent  = Cm(1.0)
    p_r2.paragraph_format.space_before = Pt(2)
    p_r2.paragraph_format.space_after  = Pt(8)
    run_label2 = p_r2.add_run('Respuesta: ')
    run_label2.bold = True
    run_label2.font.size      = Pt(10.5)
    run_label2.font.color.rgb = RGBColor(0x22, 0x22, 0x44)
    run_resp2 = p_r2.add_run(a_es)
    run_resp2.italic = True
    run_resp2.font.size      = Pt(10.5)
    run_resp2.font.color.rgb = RGBColor(0x11, 0x11, 0x44)

add_separator(doc)

# ══════════════════════════════════════════════════════════════════════════════
# CIERRE FORMAL DEL WEBINARIO
# ══════════════════════════════════════════════════════════════════════════════
add_heading1(doc, '🎬 CIERRE FORMAL DEL WEBINARIO')

add_script(doc,
    '"Thank you all so much for your incredible questions and for your energy today. It\'s '
    'conversations like these that remind me why I love this business.\n\n'
    'The recording of today\'s session will be sent to your email within 24 hours. Share it '
    'with colleagues who couldn\'t make it — this opportunity is open to every Realtor in '
    'your network.\n\n'
    'Don\'t forget: scan the QR code, register on the portal, download your Marketing Kit. '
    'Your first referral could be closer than you think.\n\n'
    'And please — reach out to us directly. Our WhatsApp is on screen. Our email is on screen. '
    'We are real people, in a real office, in a real beautiful country — and we genuinely look '
    'forward to working with you.\n\n'
    'Have a wonderful day, wherever you are. Bienvenidos al equipo. Welcome to the team."',
    speaker='GUIÓN EN INGLÉS (Agente 1 — cierre final):')

add_script_es(doc,
    '"Muchísimas gracias a todos por vuestras increíbles preguntas y por la energía de hoy. '
    'Son conversaciones como estas las que me recuerdan por qué amo este negocio.\n\n'
    'La grabación de la sesión de hoy os llegará por email en menos de 24 horas. Compartidla '
    'con colegas que no hayan podido asistir — esta oportunidad está abierta a cada Realtor '
    'de vuestra red.\n\n'
    'No lo olvidéis: escaneád el código QR, registraos en el portal, descargad vuestro '
    'Kit de Marketing. Vuestra primera referencia puede estar más cerca de lo que creéis.\n\n'
    'Y por favor — contactadnos directamente. Nuestro WhatsApp está en pantalla. Nuestro '
    'email está en pantalla. Somos personas reales, en una oficina real, en un país '
    'extraordinariamente bello — y tenemos genuinas ganas de trabajar con vosotros.\n\n'
    'Que tengáis un día maravilloso, dondequiera que estéis. Bienvenidos al equipo. '
    'Welcome to the team."',
    speaker='GUIÓN EN ESPAÑOL (Agente 1 — cierre final):')

# ══════════════════════════════════════════════════════════════════════════════
# POST-WEBINARIO
# ══════════════════════════════════════════════════════════════════════════════
add_separator(doc)
add_paragraph(doc, '📌  POST-WEBINARIO — ACCIONES INMEDIATAS (en las siguientes 2 horas):',
              bold=True, color=DKGRN, size=11, space_before=10)

acciones = [
    'Enviar email de seguimiento con la grabación y el Kit de Marketing.',
    'Contactar vía WhatsApp personal a los asistentes que dejaron su nombre de cliente en el chat.',
    'Publicar en LinkedIn/Instagram: foto de equipo + "Thank you to everyone who joined us today" + link de registro.',
    'Programar el próximo webinario para 3 semanas después y abrir inscripciones inmediatamente.',
]
for n, accion in enumerate(acciones, start=1):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(accion)
    run.font.size      = Pt(10.5)
    run.font.color.rgb = BLACK

add_separator(doc)

# ── Nota final ─────────────────────────────────────────────────────────────────
p_nota = doc.add_paragraph()
p_nota.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_nota.paragraph_format.space_before = Pt(16)
run_n1 = p_nota.add_run(
    'Documento preparado por RE/MAX Inmomás International | Uso interno del equipo de presentación\n'
)
run_n1.italic = True
run_n1.font.size      = Pt(9)
run_n1.font.color.rgb = GREY
run_n2 = p_nota.add_run('Versión 2.2 — Guión bilingüe completo EN / ES · Sin partner Smbiotica')
run_n2.italic = True
run_n2.font.size      = Pt(9)
run_n2.font.color.rgb = GREY

# ══════════════════════════════════════════════════════════════════════════════
# GUARDAR
# ══════════════════════════════════════════════════════════════════════════════
output_path = (
    '/Users/maiahonczaryk/Desktop/Proyecto Internacional/'
    'USA - Spain /Presentación B2B/Guion_Webinario_Beyond_Borders.docx'
)
doc.save(output_path)
print(f'✅  Documento guardado en:\n    {output_path}')
