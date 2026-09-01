#!/usr/bin/env python3
"""
Generate an editable Word (.docx) Referral Agreement:
  Page 1: Master Collaboration Agreement (general 25% for all referrals)
  Page 2: Annex — Individual Client Referral Form
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Colours ──
REMAX_BLUE = RGBColor(0, 61, 165)
REMAX_RED  = RGBColor(225, 27, 34)
CHARCOAL   = RGBColor(33, 37, 41)
GREY       = RGBColor(108, 117, 125)

# ── Paths ──
SCRIPT_DIR  = os.path.dirname(os.path.abspath(__file__))
LOGO_PATH   = os.path.join(SCRIPT_DIR, "remax_logo.png")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "Referral_Agreement_Template.docx")

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin   = Cm(2)
    section.right_margin  = Cm(2)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.font.color.rgb = CHARCOAL


# ── Helpers ──

def remove_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(
                '<w:tcBorders %s>'
                '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '</w:tcBorders>' % nsdecls('w')
            )
            tcPr.append(tcBorders)


def add_logo_header(title_line1, title_line2):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Inches(2.5)
    table.columns[1].width = Inches(4.0)

    cell_left = table.cell(0, 0)
    cell_left.vertical_alignment = 1
    p = cell_left.paragraphs[0]
    if os.path.exists(LOGO_PATH):
        run = p.add_run()
        run.add_picture(LOGO_PATH, width=Inches(1.6))
    else:
        run = p.add_run("RE/MAX Inmomás")
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = REMAX_BLUE

    cell_right = table.cell(0, 1)
    cell_right.vertical_alignment = 1
    p1 = cell_right.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = p1.add_run(title_line1)
    r1.bold = True
    r1.font.size = Pt(14)
    r1.font.color.rgb = CHARCOAL

    p2 = cell_right.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run(title_line2)
    r2.font.size = Pt(9)
    r2.font.color.rgb = GREY

    remove_borders(table)


def add_blue_line():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr %s>'
        '<w:bottom w:val="double" w:sz="6" w:space="1" w:color="003DA5"/>'
        '</w:pBdr>' % nsdecls('w')
    )
    pPr.append(pBdr)


def add_section_title(en_text, es_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    r_en = p.add_run(en_text)
    r_en.bold = True
    r_en.font.size = Pt(11)
    r_en.font.color.rgb = REMAX_BLUE
    r_sep = p.add_run(" / ")
    r_sep.font.size = Pt(11)
    r_sep.font.color.rgb = GREY
    r_es = p.add_run(es_text)
    r_es.bold = True
    r_es.font.size = Pt(11)
    r_es.font.color.rgb = REMAX_RED


def create_field_table():
    table = doc.add_table(rows=0, cols=2)
    table.columns[0].width = Inches(2.2)
    table.columns[1].width = Inches(4.3)
    return table


def add_field_row(table, label, value=""):
    row = table.add_row()
    c0 = row.cells[0]
    c1 = row.cells[1]
    p0 = c0.paragraphs[0]
    r0 = p0.add_run(label)
    r0.bold = True
    r0.font.size = Pt(9)
    r0.font.color.rgb = CHARCOAL
    p1 = c1.paragraphs[0]
    r1 = p1.add_run(value)
    r1.font.size = Pt(9)
    r1.font.color.rgb = CHARCOAL


def shade_table(table, color_hex="F8F9FA"):
    for row in table.rows:
        for cell in row.cells:
            shading = parse_xml(
                '<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex)
            )
            cell._tc.get_or_add_tcPr().append(shading)


def add_clause(num, title_en, title_es, text_en, text_es):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(f"{num}. {title_en} / {title_es}")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = REMAX_BLUE

    p_en = doc.add_paragraph()
    p_en.paragraph_format.space_after = Pt(4)
    p_en.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_en = p_en.add_run(text_en)
    r_en.font.size = Pt(9.5)
    r_en.font.color.rgb = CHARCOAL

    p_es = doc.add_paragraph()
    p_es.paragraph_format.space_after = Pt(6)
    p_es.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_es = p_es.add_run(text_es)
    r_es.font.size = Pt(9)
    r_es.font.color.rgb = GREY
    r_es.italic = True


def add_signature_block(title, name="", role=""):
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(14)
    r_t = p_title.add_run(title)
    r_t.bold = True
    r_t.font.size = Pt(9)
    r_t.font.color.rgb = REMAX_BLUE

    p_sig = doc.add_paragraph()
    p_sig.paragraph_format.space_before = Pt(24)
    pPr = p_sig._p.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr %s>'
        '<w:bottom w:val="single" w:sz="4" w:space="1" w:color="212529"/>'
        '</w:pBdr>' % nsdecls('w')
    )
    pPr.append(pBdr)

    fields = [
        ("Signature / Firma:", ""),
        ("Print Name / Nombre:", name),
        ("Title / Cargo:", role),
        ("Date / Fecha:", ""),
    ]
    for label, val in fields:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r_l = p.add_run(f"{label}  ")
        r_l.font.size = Pt(9)
        r_l.font.color.rgb = GREY
        if val:
            r_v = p.add_run(val)
            r_v.bold = True
            r_v.font.size = Pt(9)
            r_v.font.color.rgb = CHARCOAL


def add_footer_text(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr %s>'
        '<w:top w:val="single" w:sz="4" w:space="4" w:color="DEE2E6"/>'
        '</w:pBdr>' % nsdecls('w')
    )
    pPr.append(pBdr)
    r = p.add_run(text)
    r.font.size = Pt(7.5)
    r.font.color.rgb = GREY


# ══════════════════════════════════════════════
#  PAGE 1 — MASTER COLLABORATION AGREEMENT
# ══════════════════════════════════════════════

add_logo_header(
    "Master Referral Agreement",
    "Acuerdo Marco de Colaboración y Referidos"
)
add_blue_line()

p_main = doc.add_paragraph()
p_main.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_main.paragraph_format.space_after = Pt(2)
r1 = p_main.add_run("INTERNATIONAL REFERRAL & COLLABORATION AGREEMENT\n")
r1.bold = True; r1.font.size = Pt(13); r1.font.color.rgb = REMAX_BLUE
r2 = p_main.add_run("ACUERDO INTERNACIONAL DE REFERIDO Y COLABORACIÓN")
r2.bold = True; r2.font.size = Pt(12); r2.font.color.rgb = REMAX_RED

# Broker A
add_section_title("BROKER A — Referring Brokerage (US / Canada)", "CORREDOR A — Agencia Referente")
t1 = create_field_table()
add_field_row(t1, "Brokerage Name / Agencia:")
add_field_row(t1, "Broker of Record / Corredor:")
add_field_row(t1, "Office Address / Dirección:")
add_field_row(t1, "Broker License # / Licencia:")
add_field_row(t1, "Contact Email / Correo:")
add_field_row(t1, "Contact Phone / Teléfono:")
shade_table(t1)

# Broker B
add_section_title("BROKER B — Receiving Brokerage (Spain)", "CORREDOR B — Agencia Receptora")
t2 = create_field_table()
add_field_row(t2, "Brokerage Name / Agencia:", "Megagestión Servicios Inmobiliarios")
add_field_row(t2, "Broker of Record / Corredor:", "José Martínez Sánchez")
add_field_row(t2, "Office Address / Dirección:", "Carrer Conrado del Campo 16, 03204, Elche (Alicante), Spain")
add_field_row(t2, "CIF (Tax ID) / NIF:", "B-54829767")
add_field_row(t2, "Contact Email / Correo:", "jose.martinez@remax.es")
add_field_row(t2, "Office Phone / Teléfono:", "+34 966 665 651")
shade_table(t2)

# Clauses
add_clause("1", "Purpose & Scope", "Objeto y Alcance",
    "The purpose of this Master Agreement is to establish the general terms under which Broker A "
    "and its licensed real estate agents may refer prospective international buyers to Broker B "
    "for the purchase of real estate in Spain, specifically within the provinces of Alicante and "
    "the Costa Blanca region. This agreement covers all future referrals made by Broker A's agents "
    "during the term of this agreement. Individual client referrals shall be documented separately "
    "in the attached Annex.",
    "El objeto de este Acuerdo Marco es establecer las condiciones generales bajo las cuales el "
    "Corredor A y sus agentes inmobiliarios licenciados podrán referir compradores internacionales "
    "potenciales al Corredor B para la compra de inmuebles en España, específicamente en la provincia "
    "de Alicante y la zona de la Costa Blanca. Este acuerdo cubre todos los referidos futuros "
    "realizados por los agentes del Corredor A durante la vigencia del presente acuerdo. Cada referido "
    "individual se documentará por separado en el Anexo adjunto."
)

add_clause("2", "B2B Relationship", "Relación Comercial B2B",
    "This agreement represents a business-to-business (B2B) marketing and referral contract. "
    "Broker A is a legally licensed brokerage in its jurisdiction of origin (United States or Canada) "
    "and Broker B is a licensed real estate entity in Spain. As this is an international service "
    "referral, Broker A and its agents do not require licensing or registration in Spain to receive "
    "this B2B service fee.",
    "Este acuerdo representa un contrato de marketing y referido de empresa a empresa (B2B). "
    "El Corredor A es una agencia legalmente autorizada en su jurisdicción de origen (EE. UU. o Canadá) "
    "y el Corredor B es una entidad inmobiliaria con licencia en España. Al tratarse de un servicio de "
    "referido internacional, el Corredor A y sus agentes no requieren licencia ni registro en España "
    "para percibir estos honorarios."
)

add_clause("3", "Commission Structure", "Estructura de Comisión",
    "For every referred client whose property acquisition in Spain is successfully closed, Broker B "
    "shall pay Broker A a referral fee of 25% (twenty-five percent) of the gross buyer-side commission "
    "received by Broker B on that transaction. This rate applies uniformly to all referrals made under "
    "this agreement unless modified in a specific Annex.",
    "Por cada cliente referido cuya adquisición de propiedad en España se cierre exitosamente, el "
    "Corredor B pagará al Corredor A una comisión de referido del 25% (veinticinco por ciento) de la "
    "comisión bruta del lado comprador cobrada por el Corredor B en dicha transacción. Este porcentaje "
    "se aplica de manera uniforme a todos los referidos realizados bajo este acuerdo, salvo que se "
    "modifique en un Anexo específico."
)

# Example box
p_ex = doc.add_paragraph()
p_ex.paragraph_format.space_before = Pt(4)
p_ex.paragraph_format.space_after  = Pt(8)
p_ex.paragraph_format.left_indent  = Cm(0.5)
pPr = p_ex._p.get_or_add_pPr()
pBdr = parse_xml(
    '<w:pBdr %s>'
    '<w:left w:val="single" w:sz="12" w:space="4" w:color="003DA5"/>'
    '</w:pBdr>' % nsdecls('w')
)
pPr.append(pBdr)
r_h = p_ex.add_run("💡 Example / Ejemplo (25% Split):\n")
r_h.bold = True; r_h.font.size = Pt(9)
r_b = p_ex.add_run(
    "• Purchase Price / Precio: €400,000\n"
    "• Gross buyer commission (5%) / Comisión bruta: €20,000\n"
    "• Referral fee to Broker A (25%) / Pago al Corredor A: €5,000 EUR"
)
r_b.font.size = Pt(9); r_b.font.color.rgb = GREY

add_clause("4", "Broker Compensation", "Compensación entre Corredores",
    "All referral fees are paid strictly Broker-to-Broker. Broker A is solely responsible for "
    "any internal distribution of the commission within its brokerage. Broker B is responsible "
    "for any internal distribution within its brokerage.",
    "Todos los honorarios de referido se pagan estrictamente de Corredor a Corredor. El Corredor A "
    "es el único responsable de cualquier distribución interna de la comisión dentro de su agencia. El "
    "Corredor B es responsable de cualquier distribución interna dentro de su agencia."
)

add_clause("5", "Client Lock Period", "Vigencia del Referido",
    "Each individual referral documented in an Annex is protected for a period of 24 months from "
    "the date of registration. Broker A is entitled to the referral fee for any transaction "
    "completed by that referred client in Spain during this period.",
    "Cada referido individual documentado en un Anexo estará protegido durante un periodo de 24 meses "
    "desde la fecha de registro. El Corredor A tendrá derecho a la comisión de referido por cualquier "
    "transacción completada por dicho cliente en España durante este periodo."
)

add_clause("6", "Settlement & Tax Compliance", "Liquidación y Cumplimiento Fiscal",
    "Referral fees shall be wired to Broker A within 30 business days of transaction closing and "
    "receipt of clear funds by Broker B. US-based brokerages must provide a completed IRS Form "
    "W-8BEN-E (or Canadian equivalent) and a commercial invoice prior to payment to prevent local "
    "Spanish withholding tax.",
    "Los honorarios de referido se transferirán al Corredor A dentro de los 30 días hábiles posteriores "
    "al cierre de la transacción y cobro de los fondos por el Corredor B. Las agencias de EE. UU. "
    "deben entregar el formulario IRS W-8BEN-E (o equivalente canadiense) y una factura comercial "
    "antes del pago para evitar retenciones fiscales en España."
)

add_clause("7", "Duration & Termination", "Vigencia y Resolución",
    "This Agreement is effective for an initial period of 12 months from the date of signature, "
    "automatically renewing for successive 12-month periods unless either party provides 60 days' "
    "written notice of termination. Termination shall not affect referral rights on clients "
    "documented in existing Annexes.",
    "Este Acuerdo tendrá una vigencia inicial de 12 meses desde la fecha de firma, renovándose "
    "automáticamente por períodos sucesivos de 12 meses salvo que cualquiera de las partes comunique "
    "su resolución con un preaviso de 60 días por escrito. La resolución no afectará los derechos de "
    "referido sobre clientes documentados en Anexos existentes."
)

add_clause("8", "Governing Law", "Ley Aplicable",
    "This agreement is governed by the laws of Spain. Any disputes shall be submitted exclusively "
    "to the Courts of Alicante, Spain.",
    "Este acuerdo se rige por las leyes de España. Cualquier controversia se someterá exclusivamente "
    "a la jurisdicción de los Tribunales de Alicante, España."
)

# Signatures
add_section_title("SIGNATURES", "FIRMAS")
add_signature_block("BROKER A — Referring Brokerage / Representante Corredor A")
add_signature_block(
    "BROKER B — Megagestión / Representante Corredor B",
    name="José Martínez Sánchez",
    role="Managing Broker / Corredor de la Oficina"
)

add_footer_text(
    "Megagestión Servicios Inmobiliarios • Carrer Conrado del Campo 16, 03204, Elche (Alicante), Spain\n"
    "CIF: B-54829767 • jose.martinez@remax.es • +34 966 665 651\n"
    "© 2026 Megagestión Servicios Inmobiliarios. All Rights Reserved."
)

# ══════════════════════════════════════════════
#  PAGE 2 — ANNEX: INDIVIDUAL CLIENT REFERRAL (compact – fits on 1 page)
# ══════════════════════════════════════════════

doc.add_page_break()

# Compact header
table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.columns[0].width = Inches(2.2)
table.columns[1].width = Inches(4.3)
cl = table.cell(0, 0)
cl.vertical_alignment = 1
p = cl.paragraphs[0]
if os.path.exists(LOGO_PATH):
    run = p.add_run()
    run.add_picture(LOGO_PATH, width=Inches(1.3))
else:
    run = p.add_run("RE/MAX Inmomás")
    run.bold = True; run.font.size = Pt(14); run.font.color.rgb = REMAX_BLUE

cr = table.cell(0, 1)
cr.vertical_alignment = 1
p1 = cr.paragraphs[0]
p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r1 = p1.add_run("ANNEX / ANEXO")
r1.bold = True; r1.font.size = Pt(12); r1.font.color.rgb = REMAX_BLUE
p2 = cr.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r2 = p2.add_run("Individual Client Referral / Ficha de Referido Individual")
r2.font.size = Pt(8); r2.font.color.rgb = GREY
remove_borders(table)

# Thin blue line
p_line = doc.add_paragraph()
p_line.paragraph_format.space_before = Pt(1)
p_line.paragraph_format.space_after  = Pt(4)
pPr = p_line._p.get_or_add_pPr()
pBdr = parse_xml(
    '<w:pBdr %s>'
    '<w:bottom w:val="single" w:sz="4" w:space="1" w:color="003DA5"/>'
    '</w:pBdr>' % nsdecls('w')
)
pPr.append(pBdr)

# Ref number + date + master agreement on single line
p_ref = doc.add_paragraph()
p_ref.paragraph_format.space_after = Pt(3)
r_ref = p_ref.add_run(
    "Referral # / Nº: ______________    Date / Fecha: ______________    "
    "Master Agreement date / Fecha Acuerdo Marco: ____/____/________"
)
r_ref.font.size = Pt(8)
r_ref.font.color.rgb = REMAX_BLUE
r_ref.bold = True

# Short note
p_note = doc.add_paragraph()
p_note.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p_note.paragraph_format.space_after = Pt(5)
r_note = p_note.add_run(
    "This Annex is subject to the Master Referral Agreement. The referral is protected for 24 months. / "
    "Este Anexo está sujeto al Acuerdo Marco. El referido queda protegido durante 24 meses."
)
r_note.font.size = Pt(7.5)
r_note.font.color.rgb = GREY
r_note.italic = True


def add_compact_section(title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(title)
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = REMAX_BLUE


def create_compact_table():
    t = doc.add_table(rows=0, cols=2)
    t.columns[0].width = Inches(1.8)
    t.columns[1].width = Inches(4.7)
    return t


def add_compact_row(t, label, value=""):
    row = t.add_row()
    p0 = row.cells[0].paragraphs[0]
    p0.paragraph_format.space_after = Pt(0)
    p0.paragraph_format.space_before = Pt(0)
    r0 = p0.add_run(label)
    r0.bold = True; r0.font.size = Pt(8); r0.font.color.rgb = CHARCOAL

    p1 = row.cells[1].paragraphs[0]
    p1.paragraph_format.space_after = Pt(0)
    p1.paragraph_format.space_before = Pt(0)
    r1 = p1.add_run(value)
    r1.font.size = Pt(8); r1.font.color.rgb = CHARCOAL


# Prospect
add_compact_section("Prospect / Cliente Referido")
t5 = create_compact_table()
add_compact_row(t5, "Full Name / Nombre:")
add_compact_row(t5, "Email / Correo:")
add_compact_row(t5, "Phone / Teléfono:")
add_compact_row(t5, "Region / Región:", "Alicante Province / Costa Blanca")
add_compact_row(t5, "Budget / Presupuesto:")
add_compact_row(t5, "Property Type / Tipo:")
add_compact_row(t5, "Notes / Notas:")
shade_table(t5)

# Referral fee — inline
p_fee = doc.add_paragraph()
p_fee.paragraph_format.space_before = Pt(6)
p_fee.paragraph_format.space_after  = Pt(3)
r_ft = p_fee.add_run("Referral Fee / Comisión: ")
r_ft.bold = True; r_ft.font.size = Pt(8.5); r_ft.font.color.rgb = REMAX_BLUE
r_fv = p_fee.add_run("As per Master Agreement: 25% / Según Acuerdo Marco: 25%   ")
r_fv.font.size = Pt(8.5)
r_fm = p_fee.add_run("☐ Modified / Modificado: ________ %")
r_fm.font.size = Pt(8); r_fm.font.color.rgb = GREY

# Compact signature blocks — side by side in a table
sig_table = doc.add_table(rows=4, cols=2)
sig_table.columns[0].width = Inches(3.25)
sig_table.columns[1].width = Inches(3.25)

# Headers
for i, title in enumerate([
    "Broker A Representative / Representante Corredor A",
    "Broker B Representative / Representante Corredor B",
]):
    p = sig_table.cell(0, i).paragraphs[0]
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(0)
    r = p.add_run(title)
    r.bold = True; r.font.size = Pt(7.5); r.font.color.rgb = REMAX_BLUE

# Signature lines
for i in range(2):
    p = sig_table.cell(1, i).paragraphs[0]
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr %s>'
        '<w:bottom w:val="single" w:sz="4" w:space="1" w:color="212529"/>'
        '</w:pBdr>' % nsdecls('w')
    )
    pPr.append(pBdr)

# Name fields
for i, name_val in enumerate(["", "José Martínez Sánchez"]):
    p = sig_table.cell(2, i).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Name / Nombre: ")
    r.font.size = Pt(8); r.font.color.rgb = GREY
    if name_val:
        rv = p.add_run(name_val)
        rv.bold = True; rv.font.size = Pt(8); rv.font.color.rgb = CHARCOAL

# Date fields
for i in range(2):
    p = sig_table.cell(3, i).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Date / Fecha: ____________________")
    r.font.size = Pt(8); r.font.color.rgb = GREY

remove_borders(sig_table)

# Compact footer
p_foot = doc.add_paragraph()
p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_foot.paragraph_format.space_before = Pt(10)
pPr = p_foot._p.get_or_add_pPr()
pBdr = parse_xml(
    '<w:pBdr %s>'
    '<w:top w:val="single" w:sz="4" w:space="3" w:color="DEE2E6"/>'
    '</w:pBdr>' % nsdecls('w')
)
pPr.append(pBdr)
rf = p_foot.add_run(
    "Megagestión Servicios Inmobiliarios • CIF: B-54829767 • jose.martinez@remax.es • +34 966 665 651\n"
    "This Annex is an integral part of the Master Referral Agreement / "
    "Este Anexo forma parte integral del Acuerdo Marco de Referidos."
)
rf.font.size = Pt(7); rf.font.color.rgb = GREY

# ── Save ──
doc.save(OUTPUT_PATH)
print(f"✅ Word document saved to: {OUTPUT_PATH}")
